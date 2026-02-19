#!/usr/bin/env python3
"""
Generate CapNet Investor & Collaborator Overview (.docx)
Run: /tmp/docgen/bin/python generate_investor_doc.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
from datetime import date

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "CapNet_Investor_Overview.docx")

# -- Color palette --
DARK_NAVY = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT_BLUE = RGBColor(0x00, 0x7A, 0xCC)
ACCENT_TEAL = RGBColor(0x00, 0x96, 0x88)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY_BG = "F5F7FA"
TABLE_HEADER_BG = "1A1A2E"
TABLE_ALT_BG = "F0F4F8"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, cells_data, is_header=False, alt_row=False):
    if is_header:
        row = table.rows[0]
    else:
        row = table.add_row()

    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)

        if is_header:
            run.font.bold = True
            run.font.color.rgb = WHITE
            set_cell_shading(cell, TABLE_HEADER_BG)
        else:
            run.font.color.rgb = DARK_GRAY
            if alt_row:
                set_cell_shading(cell, TABLE_ALT_BG)

    return row


def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for level, (size, color, bold) in {
        1: (24, DARK_NAVY, True),
        2: (18, ACCENT_BLUE, True),
        3: (14, DARK_NAVY, True),
    }.items():
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = bold
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(6)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        "</w:pBdr>"
    )
    pPr.append(pBdr)


def add_callout(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_GRAY_BG}"/>')
    pPr.append(shading)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="007ACC"/>'
        "</w:pBdr>"
    )
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pf_indent = p.paragraph_format
    pf_indent.left_indent = Cm(0.5)

    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = ACCENT_BLUE

    run_t = p.add_run(text)
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = DARK_GRAY


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = DARK_GRAY
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_bold_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    return p


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    setup_styles(doc)

    # =========================================================================
    # COVER SECTION
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CapNet")
    run.font.size = Pt(42)
    run.font.color.rgb = DARK_NAVY
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("The Capability Layer for AI Agents")
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT_BLUE

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Leash, not master keys.")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = MEDIUM_GRAY

    for _ in range(4):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Investor & Collaborator Overview\n{date.today().strftime('%B %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = MEDIUM_GRAY

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run("CONFIDENTIAL")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.bold = True

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    doc.add_heading("Table of Contents", level=1)

    toc_items = [
        "1.  Executive Summary",
        "2.  The Problem \u2014 Why This Matters Now",
        "3.  The Solution \u2014 How CapNet Works",
        "4.  Technology Deep-Dive",
        "5.  Market Opportunity",
        "6.  Revenue Model & Go-to-Market Strategy",
        "7.  Competitive Landscape",
        "8.  Product Roadmap",
        "9.  The Team & The Ask",
        "10. Appendix",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.color.rgb = DARK_NAVY

    doc.add_page_break()

    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading("1. Executive Summary", level=1)

    add_callout(
        doc,
        "CapNet is a new permission layer for the Agent Era. We replace master keys with "
        "bounded, revocable, auditable capabilities enforced at the boundary where actions occur. "
        "Users set simple templates; the system compiles them into enforceable permissions. "
        "Even a compromised agent cannot exceed the leash.",
    )

    doc.add_heading("The Opportunity", level=2)
    add_body(
        doc,
        "AI agents are becoming a new class of actor on the internet. They turn natural language "
        "into real-world side effects \u2014 spending money, accessing accounts, deploying code, "
        "sending messages. They operate at machine speed and machine scale.",
    )
    add_body(
        doc,
        "Today, to let an agent act on your behalf, you must hand over raw credentials: passwords, "
        "API keys, session cookies. These are master keys with no guardrails. There is no standard "
        "way to scope what an agent can do, revoke its access instantly, or prove what it did.",
    )

    add_bold_body(doc, "CapNet solves this.")

    add_body(
        doc,
        "We are building the authorization primitive for machine actors \u2014 the trust and permission "
        "fabric that makes AI agents governable. This is not a product. It is a new fundamental layer, "
        "comparable to how TCP/IP standardized data transport or how OAuth standardized identity delegation.",
    )

    doc.add_heading("Key Principles", level=2)
    add_bullet(doc, "We don't give agents credentials. ", "We mint capabilities \u2014 ")
    add_bullet(doc, "bounded scope, time, vendors, and budget.", "")
    add_bullet(
        doc,
        "Every risky action routes through a policy enforcement boundary. ",
        "The agent never sees raw credentials. ",
    )
    add_bullet(
        doc,
        "If something goes wrong, we can prove what happened and kill it instantly. ",
        "Full audit trail, immediate revocation. ",
    )

    doc.add_heading("Current Status", level=2)
    add_body(
        doc,
        "Phase 0 is complete. We have a working end-to-end demonstration: a Chrome extension wallet UI, "
        "a local enforcement proxy, a merchant sandbox, and an SDK \u2014 all proving that scoped, "
        "revocable, cryptographically signed capabilities work today, without requiring any external "
        "partnerships or merchant changes.",
    )

    doc.add_page_break()

    # =========================================================================
    # 2. THE PROBLEM
    # =========================================================================
    doc.add_heading("2. The Problem \u2014 Why This Matters Now", level=1)

    doc.add_heading("Agents Are Becoming Actors", level=2)
    add_body(
        doc,
        "In the next 12\u201324 months, AI agents will be deployed at scale to perform real-world tasks: "
        "purchasing goods, managing calendars, booking travel, operating cloud infrastructure, accessing "
        "enterprise data, and executing financial transactions. These agents will operate autonomously, "
        "at machine speed, with real consequences.",
    )

    doc.add_heading("The Intent-to-Action Distance Is Collapsing", level=2)
    add_body(
        doc,
        "Historically, there was always a human in the loop between an intention (\"buy groceries\") "
        "and the action (clicking \"checkout\" and entering a credit card). That gap is where safety "
        "lived. AI agents collapse this distance to near zero. The new choke point is authority: "
        "who is allowed to do what, with what scope, under what auditability.",
    )

    doc.add_heading("Current Approaches Are Broken", level=2)

    # Table: current approaches
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    add_table_row(table, ["Approach", "What It Does", "Why It Fails for Agents"], is_header=True)
    rows = [
        ("Shared Credentials", "Give agent your password/API key", "Master key with no scope limits; can't revoke without rotating; no audit trail"),
        ("OAuth Scopes", "Grant broad read/write access tokens", "Designed for identity, not scoped authority; no budget limits, category blocks, or instant revocation"),
        ("API Gateways", "Rate limiting and authentication", "Rate limiting is not policy enforcement; no concept of budget, vendor, or time-bounded authority"),
        ("Manual Approval", "Human approves every action", "Defeats the purpose of autonomous agents; doesn't scale"),
    ]
    for i, row in enumerate(rows):
        add_table_row(table, row, alt_row=(i % 2 == 1))

    add_body(doc, "")

    doc.add_heading("The Consequences Are Real", level=2)
    add_bullet(doc, "Unauthorized spending and financial overreach by compromised or buggy agents")
    add_bullet(doc, "Data exfiltration through overly broad permissions")
    add_bullet(doc, "Regulatory and compliance exposure with no audit trail")
    add_bullet(doc, "Enterprise liability for agent actions taken without proper authorization")
    add_bullet(doc, "Consumer trust erosion when agents act outside expected bounds")

    add_callout(
        doc,
        "Nobody has solved governable authority for machine actors. That's CapNet.",
        "The missing layer: ",
    )

    doc.add_page_break()

    # =========================================================================
    # 3. THE SOLUTION
    # =========================================================================
    doc.add_heading("3. The Solution \u2014 How CapNet Works", level=1)

    doc.add_heading("The Capability Model", level=2)
    add_body(
        doc,
        "CapNet defines the standard \"atom\" of authority for agents. Every capability is:",
    )
    add_bullet(doc, "narrowly targeted to specific resources and actions", "Scoped \u2014 ")
    add_bullet(doc, "expires, with not-before constraints", "Time-bounded \u2014 ")
    add_bullet(doc, "can be killed instantly by the user", "Revocable \u2014 ")
    add_bullet(doc, "derived sub-capabilities can only reduce authority (never expand)", "Composable \u2014 ")
    add_bullet(doc, "full chain of who requested, who approved, what policy allowed it", "Traceable \u2014 ")

    doc.add_heading("Architecture", level=2)
    add_body(
        doc,
        "CapNet operates as a sidecar architecture with three components:",
    )

    # Architecture diagram as a simple table
    arch_table = doc.add_table(rows=1, cols=4)
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    arch_table.style = "Table Grid"
    add_table_row(arch_table, ["Component", "Role", "Description", "User Interaction"], is_header=True)
    arch_rows = [
        ("Wallet UI\n(Chrome Extension)", "Consent Surface", "Policy templates, capability issuance, revocation button, audit trail viewer", "User sets rules and monitors activity"),
        ("Enforcement Proxy", "Policy Boundary", "Validates signatures, enforces constraints (budget, vendor, category, time), emits receipts", "Invisible to user; transparent to agent"),
        ("Resource / Merchant", "Action Target", "The service the agent interacts with (store, API, SaaS tool)", "No changes required; proxy handles everything"),
        ("SDK", "Agent Interface", "Client library for agent frameworks to submit action requests through the proxy", "Developer integration point"),
    ]
    for i, row in enumerate(arch_rows):
        add_table_row(arch_table, row, alt_row=(i % 2 == 1))

    add_body(doc, "")

    add_callout(
        doc,
        "The proxy is the sole enforcement boundary. The agent never receives raw credentials. "
        "Even if an agent is fully compromised, it cannot exceed the bounds of its capability.",
        "Key principle: ",
    )

    doc.add_heading("The Demo Story \u2014 5 Steps", level=2)
    add_body(doc, "CapNet's value becomes obvious in a 2-minute demonstration:")

    steps = [
        ("1. User sets policy:", "\"Groceries, $200 max, block alcohol\" \u2014 via a simple template in the wallet UI."),
        ("2. System mints capability:", "A cryptographically signed, time-bounded, executor-bound permission is issued."),
        ("3. Agent shops:", "The agent browses a catalog and checks out. The proxy verifies constraints and allows it."),
        ("4. Agent tries forbidden item:", "The agent attempts to buy wine. Blocked instantly with a clear reason: \"Category blocked: alcohol.\""),
        ("5. User revokes:", "One click. All further agent actions are denied immediately. Full audit trail is available."),
    ]
    for bold, text in steps:
        add_bullet(doc, text, bold)

    add_body(doc, "")

    doc.add_heading("What Makes CapNet Different from OAuth / IAM", level=2)

    comp_table = doc.add_table(rows=1, cols=4)
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comp_table.style = "Table Grid"
    add_table_row(comp_table, ["Feature", "OAuth / API Keys", "IAM / RBAC", "CapNet"], is_header=True)
    comp_rows = [
        ("Primary question", "\"Who is this?\"", "\"What role does this have?\"", "\"What can this agent do right now?\""),
        ("Scope granularity", "Broad (read, write)", "Role-based", "Fine-grained (budget, vendor, category, time)"),
        ("Revocation speed", "Token expiry (minutes-hours)", "Policy propagation (minutes)", "Instant (proxy-enforced)"),
        ("Audit trail", "Access logs", "Permission change logs", "Per-action receipts with denial reasons"),
        ("Budget enforcement", "Not supported", "Not supported", "Native (max_amount_cents)"),
        ("Category blocking", "Not supported", "Not supported", "Native (blocked_categories)"),
        ("Delegation control", "No attenuation", "Role inheritance", "Monotone reduction (can only shrink)"),
        ("Agent awareness", "Not designed for agents", "Not designed for agents", "Built for agent-first workflows"),
    ]
    for i, row in enumerate(comp_rows):
        add_table_row(comp_table, row, alt_row=(i % 2 == 1))

    doc.add_page_break()

    # =========================================================================
    # 4. TECHNOLOGY DEEP-DIVE
    # =========================================================================
    doc.add_heading("4. Technology Deep-Dive", level=1)

    add_body(
        doc,
        "This section is for technical collaborators, engineers, and security-minded investors who want "
        "to understand how CapNet works under the hood.",
    )

    doc.add_heading("Cryptographic Foundation", level=2)
    add_bullet(doc, "All capabilities are signed using Ed25519 with domain separation prefixes, preventing cross-protocol signature reuse.")
    add_bullet(doc, "Canonical JSON serialization (deterministic key ordering) ensures signature stability across implementations.")
    add_bullet(doc, "Browser-safe base64 encoding works identically in Node.js and Chrome extension environments.")
    add_bullet(doc, "Key length validation (32-byte public keys, 64-byte signatures) prevents malformed input.")

    doc.add_heading("CapDoc v0.1 \u2014 The Capability Object", level=2)
    add_body(
        doc,
        "Every capability is a CapDoc \u2014 a JSON object containing:",
    )

    capdoc_table = doc.add_table(rows=1, cols=3)
    capdoc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    capdoc_table.style = "Table Grid"
    add_table_row(capdoc_table, ["Field", "Type", "Purpose"], is_header=True)
    capdoc_rows = [
        ("version", "\"capdoc/0.1\"", "Schema version for forward compatibility"),
        ("cap_id", "string", "Unique identifier (cap_<timestamp>_<rand>)"),
        ("issuer", "object", "Issuer identity and Ed25519 public key"),
        ("executor", "object", "Bound agent identity and public key"),
        ("resource", "object", "Target resource type and vendor"),
        ("actions", "array", "Permitted verbs (e.g., [\"spend\"])"),
        ("constraints", "object", "Budget, vendor allowlist, blocked categories, time window"),
        ("revocation", "object", "Revocation mode (strict/lease/one-time) and oracle pointer"),
        ("proof", "object", "Ed25519 signature over the canonical unsigned payload"),
    ]
    for i, row in enumerate(capdoc_rows):
        add_table_row(capdoc_table, row, alt_row=(i % 2 == 1))

    add_body(doc, "")

    doc.add_heading("Enforcement Pipeline", level=2)
    add_body(
        doc,
        "When an agent submits an action request, the proxy evaluates it through a strict, ordered pipeline:",
    )
    steps_pipeline = [
        ("1. Signature verification:", "Is the capability's cryptographic signature valid? Reject if tampered."),
        ("2. Executor binding:", "Does the requesting agent's public key match the capability's executor binding? Reject impostors."),
        ("3. Time semantics:", "Is the capability within its valid time window (not_before \u2264 now \u2264 expires_at)? Reject expired or premature requests."),
        ("4. Revocation check:", "Has this capability been revoked? Reject immediately if so."),
        ("5. Constraint enforcement:", "Does the action comply with budget limits, vendor allowlists, and category blocks? Reject violations with specific reasons."),
    ]
    for bold, text in steps_pipeline:
        add_bullet(doc, text, bold)

    add_body(doc, "")
    add_body(
        doc,
        "Every action \u2014 whether allowed or denied \u2014 produces a signed receipt with the decision, "
        "reason, timestamp, and all relevant identifiers. Receipts are stored in an append-only log.",
    )

    doc.add_heading("Security Model", level=2)
    add_body(doc, "CapNet's security posture follows defense-in-depth principles:")
    add_bullet(doc, "Design for containment, not just prevention. Assume agents can be buggy or malicious.", "Assume breach: ")
    add_bullet(doc, "No action is permitted unless a valid, signed capability explicitly authorizes it.", "Default deny: ")
    add_bullet(doc, "Capabilities bound to specific agents prevent stolen capabilities from being used by unauthorized actors.", "Executor binding: ")
    add_bullet(doc, "Even valid capabilities are constrained to specific budgets, vendors, categories, and time windows.", "Blast-radius containment: ")
    add_bullet(doc, "Instant revocation through the proxy \u2014 no propagation delay.", "Kill switch: ")

    doc.add_heading("Attenuation & Delegation (In Development)", level=2)
    add_body(
        doc,
        "CapNet supports delegation through monotone attenuation: a derived sub-capability can only "
        "reduce authority, never expand it. A parent capability with a $200 budget can delegate a "
        "sub-capability with $50, but never $300. Expiry can be shortened, never extended. Vendor lists "
        "can be narrowed, never broadened. This is enforced mechanically \u2014 no interpretation disputes.",
    )

    doc.add_page_break()

    # =========================================================================
    # 5. MARKET OPPORTUNITY
    # =========================================================================
    doc.add_heading("5. Market Opportunity", level=1)

    doc.add_heading("The AI Agent Economy Is Emerging", level=2)
    add_body(
        doc,
        "The AI agent market is at an inflection point. Major technology companies \u2014 OpenAI, Anthropic, "
        "Google, Microsoft, and dozens of startups \u2014 are investing heavily in agent capabilities. "
        "Agents are moving from research demos to production deployments across enterprise IT, customer "
        "service, software development, finance, and consumer applications.",
    )

    add_body(
        doc,
        "Every agent that takes a real-world action needs an authorization layer. Currently, there is no "
        "standard solution. Each team builds ad-hoc guardrails or, more commonly, simply shares credentials. "
        "This is the same state the internet was in before OAuth standardized identity delegation.",
    )

    doc.add_heading("Market Sizing", level=2)

    market_table = doc.add_table(rows=1, cols=3)
    market_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    market_table.style = "Table Grid"
    add_table_row(market_table, ["Segment", "Relevance to CapNet", "Market Signal"], is_header=True)
    market_rows = [
        ("AI Agent Platforms", "Every agent framework needs authorization primitives", "OpenAI, Anthropic, LangChain, CrewAI, AutoGen all building agent tooling"),
        ("Enterprise IAM", "Agents are a new actor class that existing IAM doesn't handle", "$18B+ market seeking agent-aware authorization"),
        ("API Security", "Agent-to-API interactions need policy enforcement beyond rate limits", "Growing demand for fine-grained API governance"),
        ("Compliance & Audit", "Regulated industries need provable audit trails for agent actions", "Financial services, healthcare, government mandating AI governance"),
        ("Developer Tools", "SDK and infrastructure for building safe agent integrations", "Fastest-growing segment in developer tooling"),
    ]
    for i, row in enumerate(market_rows):
        add_table_row(market_table, row, alt_row=(i % 2 == 1))

    add_body(doc, "")

    doc.add_heading("Why Now", level=2)
    add_bullet(doc, "Agents are shipping to production in 2025\u20132026 at scale, but authorization infrastructure hasn't kept pace.", "Timing: ")
    add_bullet(doc, "Without a safe delegation layer, credential leakage and agent overreach will produce consumer losses, enterprise breaches, and regulatory backlash.", "Urgency: ")
    add_bullet(doc, "The standard that wins early becomes the default. OAuth won because it was first with a workable spec. CapNet aims to be OAuth for agent authority.", "First-mover: ")

    add_callout(
        doc,
        "If CapNet wins the unit of authority for agents, we become fundamental infrastructure \u2014 "
        "not a feature, not a product, but a layer.",
        "The bet: ",
    )

    doc.add_page_break()

    # =========================================================================
    # 6. REVENUE MODEL & GTM
    # =========================================================================
    doc.add_heading("6. Revenue Model & Go-to-Market Strategy", level=1)

    doc.add_heading("Business Model: Open-Core Infrastructure", level=2)
    add_body(
        doc,
        "CapNet follows the proven open-core model that has built companies like HashiCorp, Elastic, "
        "Confluent, and Datadog. The core protocol and SDK are open source to maximize developer adoption "
        "and ecosystem growth. Commercial products are built on top for enterprise customers.",
    )

    doc.add_heading("Revenue Streams", level=2)

    rev_table = doc.add_table(rows=1, cols=4)
    rev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rev_table.style = "Table Grid"
    add_table_row(rev_table, ["Stream", "Description", "Pricing Model", "Timeline"], is_header=True)
    rev_rows = [
        (
            "CapNet Cloud\n(Managed Service)",
            "Hosted enforcement proxy, revocation oracle, receipt storage, and compliance dashboards. Zero infrastructure for customers.",
            "Usage-based: per-capability-issued + per-action-enforced. Free tier for developers (1,000 actions/month).",
            "Phase 1\n(6-12 months)",
        ),
        (
            "Enterprise Platform",
            "Self-hosted deployment with SSO integration, custom policy engines, multi-tenant support, SLA guarantees, and dedicated support.",
            "Annual contract: $50K-500K+ depending on scale and support tier.",
            "Phase 2\n(12-18 months)",
        ),
        (
            "Certification & Conformance",
            "\"CapNet Verified\" certification for agent frameworks and SaaS platforms. Conformance test suite licensing.",
            "Per-certification fee + annual renewal. Test suite licensing for platform partners.",
            "Phase 2\n(12-18 months)",
        ),
        (
            "Policy Marketplace",
            "Curated, audited policy templates for regulated industries (financial services, healthcare, government).",
            "Subscription per industry pack, or included in enterprise tier.",
            "Phase 3\n(18-24 months)",
        ),
    ]
    for i, row in enumerate(rev_rows):
        add_table_row(rev_table, row, alt_row=(i % 2 == 1))

    add_body(doc, "")

    doc.add_heading("Go-to-Market Strategy", level=2)

    doc.add_heading("Phase 1: Developer Adoption (Bottoms-Up)", level=3)
    add_bullet(doc, "Open-source SDK with integrations for major agent frameworks (LangChain, OpenAI Agents, Anthropic tool-use, CrewAI, AutoGen)")
    add_bullet(doc, "5-minute quickstart and live demo that any developer can run locally")
    add_bullet(doc, "Developer documentation, tutorials, and example integrations")
    add_bullet(doc, "Community building: Discord/Slack, conference talks, blog posts, open spec")
    add_bullet(doc, "Target: 1,000+ developers using the SDK within 6 months")

    doc.add_heading("Phase 2: Enterprise Pilot", level=3)
    add_bullet(doc, "Partner with 3-5 enterprises running internal AI agent workflows")
    add_bullet(doc, "Use cases: IT operations, finance approvals, code deployment, data access control")
    add_bullet(doc, "Prove: reduced blast radius vs. shared secrets, audit quality, revoke speed")
    add_bullet(doc, "Convert pilots to paid enterprise contracts")

    doc.add_heading("Phase 3: Platform Play", level=3)
    add_bullet(doc, "Agent-to-SaaS connectors (Stripe, GitHub, Google Workspace, Slack, AWS)")
    add_bullet(doc, "Cross-organization delegated trust (one org grants another org's agents scoped access)")
    add_bullet(doc, "Invite major AI labs (OpenAI, Anthropic, Google) to co-develop the spec")
    add_bullet(doc, "Position CapNet as the default authorization standard for the agent ecosystem")

    doc.add_heading("Pricing Philosophy", level=2)
    add_body(
        doc,
        "Like Stripe (\"we make money when you make money\"), CapNet's usage-based pricing aligns with "
        "customer success. The more agents a customer deploys, the more capabilities they issue, the more "
        "value CapNet provides, and the more revenue we generate. This creates natural expansion revenue "
        "within accounts.",
    )

    doc.add_page_break()

    # =========================================================================
    # 7. COMPETITIVE LANDSCAPE
    # =========================================================================
    doc.add_heading("7. Competitive Landscape", level=1)

    doc.add_heading("Why Existing Solutions Don't Solve This", level=2)

    add_body(doc, "The agent authorization problem is genuinely new. Existing solutions were designed for human users, not autonomous machine actors:")

    add_bullet(doc, "Answers \"who is this?\" not \"what can this agent do right now, and can I stop it?\" OAuth is identity. CapNet is scoped, revocable, auditable authority with enforcement at the boundary.", "OAuth / OIDC: ")
    add_bullet(doc, "Role-based access control assigns static roles. Agents need dynamic, time-bounded, budget-aware permissions that can be revoked instantly. IAM doesn't model \"spend $50 at this vendor for the next 2 hours.\"", "IAM / RBAC: ")
    add_bullet(doc, "Provide rate limiting and authentication, not policy enforcement. No concept of budget limits, category blocks, or per-action audit receipts.", "API Gateways: ")
    add_bullet(doc, "Distributed consensus adds unnecessary overhead and latency. CapNet uses cryptographic signatures (Ed25519), not blockchain. The enforcement is local and instant.", "Blockchain / Web3: ")
    add_bullet(doc, "Zero-trust assumes \"never trust, always verify\" for network access. CapNet goes further: even verified agents are constrained to specific scoped actions. It's authorization at the action level, not the network level.", "Zero Trust: ")

    doc.add_heading("CapNet's Moat", level=2)
    add_bullet(doc, "First to define the primitive and spec for agent capability-based authorization.", "First-mover on the standard: ")
    add_bullet(doc, "Open protocol + conformance suite creates lock-in through ecosystem, not vendor lock-in.", "Network effects: ")
    add_bullet(doc, "Working demo today. Not a whitepaper \u2014 running code with cryptographic enforcement.", "Running code: ")
    add_bullet(doc, "If CapNet becomes how you express agent permissions, we're the TCP/IP of agency. The moat is the standard, not the implementation.", "Standard ownership: ")

    doc.add_page_break()

    # =========================================================================
    # 8. ROADMAP
    # =========================================================================
    doc.add_heading("8. Product Roadmap", level=1)

    road_table = doc.add_table(rows=1, cols=4)
    road_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    road_table.style = "Table Grid"
    add_table_row(road_table, ["Phase", "Timeline", "Deliverables", "Success Criteria"], is_header=True)
    road_rows = [
        (
            "Phase 0\n(COMPLETE)",
            "Completed\nFeb 2026",
            "Agent Sandbox Wallet + Proxy\n\u2022 CapDoc v0.1 schema + Ed25519 crypto\n\u2022 Proxy enforcement (budget, vendor, category, time, executor)\n\u2022 Chrome extension wallet UI\n\u2022 Merchant sandbox + SDK\n\u2022 Receipts and audit trail\n\u2022 Revocation with persistence",
            "\u2022 Working end-to-end demo\n\u2022 Allow/deny/revoke cycle verified\n\u2022 Cross-platform (Win/Mac/Linux)",
        ),
        (
            "Phase 1\nReal Integration",
            "3-6 months",
            "\u2022 Stripe test-mode integration (real payment rails)\n\u2022 GitHub API integration (agent can't delete repos)\n\u2022 Delegation / attenuation (sub-capabilities)\n\u2022 Conformance test suite\n\u2022 Investor-mode demo polish",
            "\u2022 CapNet gates something real\n\u2022 \"Not a toy\" proven\n\u2022 First 3-5 integrations live",
        ),
        (
            "Phase 2\nEnterprise",
            "6-18 months",
            "\u2022 Enterprise proxy deployment (self-hosted)\n\u2022 SSO / identity provider integration\n\u2022 Policy engine with custom rules\n\u2022 Compliance dashboards + receipt export\n\u2022 Multi-tenant support",
            "\u2022 3-5 enterprise pilots\n\u2022 First paid contracts\n\u2022 Reduced blast radius vs. secrets proven",
        ),
        (
            "Phase 3\nPlatform",
            "18-36 months",
            "\u2022 Agent-to-SaaS connector marketplace\n\u2022 Cross-org delegated trust\n\u2022 Managed CapNet Cloud (hosted service)\n\u2022 Spec published for multi-vendor adoption\n\u2022 AI lab partnerships (OpenAI, Anthropic, Google)",
            "\u2022 Third parties build \"for CapNet\"\n\u2022 Industry standard adoption\n\u2022 Sustainable recurring revenue",
        ),
        (
            "North Star",
            "3-5 years",
            "Universal Capability Fabric\n\u2022 Default authorization layer for all agent interactions\n\u2022 OS-level integration\n\u2022 Hardware/TEE attestation\n\u2022 Cross-device, cross-service, cross-org trust fabric",
            "\"Of course we use CapNet for agent authorization\"",
        ),
    ]
    for i, row in enumerate(road_rows):
        add_table_row(road_table, row, alt_row=(i % 2 == 1))

    doc.add_page_break()

    # =========================================================================
    # 9. TEAM & ASK
    # =========================================================================
    doc.add_heading("9. The Team & The Ask", level=1)

    doc.add_heading("The Team", level=2)
    add_body(
        doc,
        "[Team bios to be added. Include founding team background, relevant expertise in "
        "security, infrastructure, AI/ML, and distributed systems.]",
    )

    doc.add_heading("What We're Looking For", level=2)

    doc.add_heading("Funding", level=3)
    add_body(
        doc,
        "Seed-stage investment to accelerate from working demo to first real integrations and "
        "enterprise pilots. Capital will fund:",
    )
    add_bullet(doc, "Engineering team expansion (proxy hardening, SDK integrations, enterprise features)")
    add_bullet(doc, "Developer relations and community building")
    add_bullet(doc, "First enterprise pilot program")
    add_bullet(doc, "Spec development and conformance suite")

    doc.add_heading("Strategic Partnerships", level=3)
    add_bullet(doc, "AI labs building agent frameworks (integration partnerships)")
    add_bullet(doc, "Enterprise customers willing to pilot agent authorization")
    add_bullet(doc, "SaaS platforms interested in native CapNet connectors")
    add_bullet(doc, "Security and compliance firms for co-marketing and certification")

    doc.add_heading("Technical Collaborators", level=3)
    add_bullet(doc, "Engineers with experience in authorization systems, cryptography, or distributed systems")
    add_bullet(doc, "Security researchers interested in agent safety and capability-based security")
    add_bullet(doc, "Contributors to the open protocol and conformance suite")

    doc.add_heading("The 3-5 Year Success State", level=2)
    add_body(doc, "If CapNet succeeds, people will say:")
    add_callout(doc, "\"We don't give agents raw credentials.\"")
    add_callout(doc, "\"We mint CapNet capabilities.\"")
    add_callout(doc, "\"Every risky action routes through CapNet policies.\"")
    add_callout(doc, "\"When something goes wrong, we can prove what happened and shut it down instantly.\"")
    add_body(doc, "That's internet-grade default behavior, not a feature.")

    doc.add_page_break()

    # =========================================================================
    # 10. APPENDIX
    # =========================================================================
    doc.add_heading("10. Appendix", level=1)

    doc.add_heading("A. Investor FAQ", level=2)

    faqs = [
        (
            "\"Agent spending seems niche \u2014 what makes this a big opportunity?\"",
            "Spending is the easiest way to show the primitive. The investment isn't groceries \u2014 it's the "
            "control plane for delegated authority. Today, either you give an agent raw credentials or you "
            "don't let it act. CapNet creates a third option: safe delegation with immediate revocation and "
            "provable audit. Once that exists, it becomes the default way any system accepts agent actions. "
            "The TAM is every agent-to-service interaction, across every industry.",
        ),
        (
            "\"How is this different from OAuth/IAM?\"",
            "OAuth answers \"who is this?\" CapNet answers \"what can this agent do right now, and can I "
            "stop it?\" OAuth is identity. CapNet is scoped, revocable, auditable authority with enforcement "
            "at the boundary. They're complementary, not competitive \u2014 CapNet works with existing identity "
            "systems.",
        ),
        (
            "\"Why will developers adopt this?\"",
            "Same reason they adopted HTTPS: it's the only way to do it safely. When agents routinely take "
            "real-world actions, \"give it my API key\" stops being acceptable. CapNet is the path that "
            "doesn't require trusting the agent. Plus, our bottoms-up GTM means a single engineer can adopt "
            "without committee approval \u2014 like Stripe, Twilio, or Firebase.",
        ),
        (
            "\"What's the moat?\"",
            "First-mover on the primitive + spec. If CapNet becomes how you express agent permissions, we're "
            "the TCP/IP of agency. The moat is the standard, not the implementation. Open-core ensures "
            "ecosystem growth while enterprise features and managed services capture revenue.",
        ),
        (
            "\"Won't the big AI labs just build this themselves?\"",
            "AI labs are building agent capabilities, not authorization infrastructure. Authorization is a "
            "cross-cutting concern that works across all frameworks. Labs are incentivized to adopt a standard "
            "rather than build proprietary solutions \u2014 it reduces their liability and increases trust in "
            "their platforms. We're building for interoperability, which is harder to do from inside one ecosystem.",
        ),
    ]

    for question, answer in faqs:
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(question)
        run_q.bold = True
        run_q.font.size = Pt(11)
        run_q.font.color.rgb = DARK_NAVY

        add_body(doc, answer)
        add_body(doc, "")

    doc.add_heading("B. Try It Yourself", level=2)
    add_body(doc, "The Phase 0 demo runs locally in under 5 minutes:")

    demo_steps = [
        "Prerequisites: Node.js 18.x, Chrome browser",
        "Clone the repository and run: npm install",
        "Start services: npm run dev",
        "Build and load the Chrome extension (extension/dist/) as an unpacked extension",
        "Run the demo: npm run demo",
        "Watch: allowed purchase, denied purchase (alcohol), revocation, post-revoke denial, full audit trail",
    ]
    for i, step in enumerate(demo_steps, 1):
        add_bullet(doc, step, f"{i}. ")

    doc.add_heading("C. Contact", level=2)
    add_body(doc, "[Contact information to be added]")
    add_body(doc, "[Website / repository URL to be added]")

    add_horizontal_rule(doc)

    # Footer note
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(f"CapNet \u2014 {date.today().strftime('%B %Y')} \u2014 Confidential")
    run.font.size = Pt(9)
    run.font.color.rgb = MEDIUM_GRAY
    run.italic = True

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    build_document()
